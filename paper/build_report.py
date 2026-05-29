# -*- coding: utf-8 -*-
# 生成实验报告 report_notegen.docx
# - 严格遵循模板格式：南京工程学院 人工智能课程综合大作业说明书
# - 6 章 + 参考文献 + 成绩页
# - 字体：宋体 12pt 正文；黑体 16pt 一级标题
# - 复用 PPT 的 3 张截图
# 运行：python build_report.py
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / 'template_raw.docx'
OUT = ROOT / 'report_notegen.docx'
SCREENSHOT_DIR = ROOT.parent / 'presentation' / 'screenshots'

shutil.copy(TEMPLATE, OUT)
doc = Document(str(OUT))

# 清空文档内容（保留 section / styles）
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)
for t in list(doc.tables):
    t._element.getparent().remove(t._element)


# ============================================================
# 工具函数
# ============================================================

def set_cn_font(run, font='宋体', size=12, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)


def add_para(text='', *, font='宋体', size=12, bold=False, align=None,
             first_indent=True, line_spacing=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if first_indent and text:
        pf.first_line_indent = Pt(size * 2)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        set_cn_font(r, font=font, size=size, bold=bold, color=color)
    return p


def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_cn_font(r, font='黑体', size=16, bold=True)
    return p


def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_cn_font(r, font='黑体', size=14, bold=True)
    return p


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_cn_font(r, font='黑体', size=12, bold=True)
    return p


def add_code(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(10.5)
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(24)
        r0 = p.add_run('● ')
        set_cn_font(r0, font='宋体', size=12)
        r1 = p.add_run(it)
        set_cn_font(r1, font='宋体', size=12)


def add_image(path, width_in=5.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = None
        cr = cp.add_run(caption)
        set_cn_font(cr, font='宋体', size=10.5, bold=True)


def add_image_placeholder(text='[ 此处留给同学手动插入截图 ]', height_in=2.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_cn_font(r, font='宋体', size=11, bold=True,
                color=RGBColor(0x80, 0x80, 0x80))
    for _ in range(int(height_in * 2)):
        bp = doc.add_paragraph()
        bp.paragraph_format.first_line_indent = None
        bp.paragraph_format.space_after = Pt(0)
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bp.add_run('　')
        set_cn_font(br, font='宋体', size=11)


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None


def page_break():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._element.append(br)


# ============================================================
# 封面
# ============================================================
def build_cover():
    for _ in range(3):
        add_blank()
    add_para('南京工程学院', size=22, align=WD_ALIGN_PARAGRAPH.CENTER,
             first_indent=False, line_spacing=1.5)
    add_blank()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('人工智能')
    set_cn_font(r, font='黑体', size=36, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('课程综合大作业说明书')
    set_cn_font(r, font='黑体', size=36, bold=True)
    for _ in range(4):
        add_blank()
    info_rows = [
        ('题       目', '基于深度学习的网课视频摘要与笔记生成系统'),
        ('专       业', '数字媒体技术'),
        ('班       级', '数字媒体技术212'),
        ('学 生 姓 名', '许书恺'),
        ('学       号', '202210630'),
    ]
    for label, value in info_rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = None
        r = p.add_run(f'{label}    {value}')
        set_cn_font(r, font='宋体', size=16, bold=True)
    add_blank()
    add_blank()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('2026 年  5 月')
    set_cn_font(r, font='宋体', size=16, bold=True)
    page_break()


# ============================================================
# 目录占位
# ============================================================
def build_toc():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('目  录')
    set_cn_font(r, font='黑体', size=16, bold=True)
    add_blank()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = '[请在 Word 中右键此处 → 更新域，生成完整目录]'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    run._element.append(fldChar4)
    page_break()


# ============================================================
# 第一章
# ============================================================
def chapter_1():
    add_h1('一、选题的意义和研究背景')

    add_h2('(一) 选题的意义')
    add_para('教育数字化与在线学习需求：')
    add_para('近年来，网课、慕课、教学视频成为主要的学习载体。B 站、YouTube、Coursera 等'
             '平台上的优质教学资源呈爆炸式增长。然而，一节完整的网课视频时长往往达到 1 到 2 小时，'
             '学习者看完一遍后再回头复习需要付出极大的时间成本，传统的「边看视频边手写笔记」也难以兼顾听讲与记录。'
             '如何利用人工智能技术，自动从教学视频中提取关键内容并生成结构化笔记，已成为教育数字化领域亟待解决的实际问题。')

    add_para('学习效率提升：')
    add_para('自动化结构化笔记能将原本需要重复观看的视频压缩为可快速浏览的文本，'
             '并通过章节切分、知识点高亮、术语表索引等手段提供「二次学习」的快捷入口。'
             '本系统通过端到端 pipeline 将视频转化为带目录、摘要、知识点和术语表的 Markdown 笔记，'
             '可显著降低复习成本，提升学习效率。')

    add_para('学习场景特化：')
    add_para('现有的会议纪要工具（如 Otter.ai、Notion AI）将所有视频或音频内容套用同一种摘要模板，'
             '这种「一刀切」的产品设计对教学视频并不友好。教学视频需要的是知识点速览、术语表、'
             '章末小结等学习专属元素，而 Vlog、演讲等其他类型视频又需要不同的呈现方式。'
             '本项目针对学习场景特化设计了 5 类学习元素模板，并通过启发式分类器对不同类型视频做差异化分发，'
             '解决了通用产品在学习场景下的结构空洞问题。')

    add_para('学术研究价值：')
    add_para('从学术研究的角度，本项目融合了多个深度学习子领域：基于 CTranslate2 加速的 ASR（自动语音识别）、'
             '基于 4-bit AWQ 量化的大语言模型作章节切分器（LLM-as-Segmenter）、视觉语言模型（VLM）作多模态线索补充、'
             '以及 CLIP 视觉编码器作章节边界辅助信号。系统在 32 视频的跨域 benchmark 上验证了'
             '多模态融合权重 α=0.3 是稳健甜点、LLM 切分配合三层兜底机制可达 100% 覆盖率等具有方法论价值的实证结论，'
             '对后续视频结构化研究具有参考意义。')

    add_h2('(二) 研究背景')
    add_para('经济背景：')
    add_para('中国在线教育市场规模持续扩大，教育部统计显示 2024 年高校在线开放课程访问量已突破 50 亿人次。'
             'B 站学习类视频日均播放量超过 6 亿次。在如此庞大的教学视频内容池中，结构化检索与复习工具的'
             '市场空缺巨大，自动笔记生成系统具有明确的商业应用价值。')

    add_para('技术背景：')
    add_para('近几年深度学习技术在多个相关方向上取得突破性进展：OpenAI Whisper 系列模型将多语种 ASR 的'
             '错误率降低到接近人类水平；以 Qwen2.5、Llama-3、DeepSeek-V3 为代表的开源大语言模型在中文'
             '长文本理解上达到了可工程化部署的成熟度；CLIP 和 Chinese-CLIP 提供了跨模态语义对齐的视觉编码器；'
             'AWQ、GPTQ 等 4-bit 量化技术让 7B 级模型可在消费级 GPU 上单卡运行。'
             '这些底层技术的成熟，为本项目在单台个人电脑上实现完整教学视频结构化 pipeline 提供了可能。')

    add_para('社会背景：')
    add_para('移动互联网时代用户呈现碎片化学习的特征，但碎片化的输入与结构化的输出之间存在巨大鸿沟。'
             '学生希望在通勤、午休等碎片时间快速回顾一节课的核心内容，传统的「重新看视频」无法满足这一需求。'
             '本系统的章节级目录、知识点速览和术语表索引正是面向碎片化复习场景设计，'
             '用户可以在 3 分钟内回顾一节 1 小时课程的核心要点。')

    add_para('学术背景：')
    add_para('现有视频摘要研究多面向短视频，长教学视频缺少专门方案。'
             'TextTiling、C99 等无监督文本分段方法可以处理 ASR 转写，但对 PPT 教学视频中 slide 翻页频繁'
             '却话题未变这一特殊场景效果不佳。基于 GPT-4 的黑盒切分研究虽有探索但缺少失败模式与'
             '回退策略的系统讨论。本项目通过本地 Qwen2.5-7B-AWQ 配合 B1 两步法、retry-with-feedback 与'
             '程序化 repair 三层兜底，在 24 视频 corpus 上达到 100% 切分覆盖率，'
             '对 LLM-as-Segmenter 方向贡献了可复现的工程方案与实证数据。')

    add_para('综上所述，选题「基于深度学习的网课视频摘要与笔记生成系统」既有明确的现实意义和应用价值，'
             '也具备深厚的研究背景和学术意义。')
    add_blank()


# ============================================================
# 第二章
# ============================================================
def chapter_2():
    add_h1('二、相关技术简介和流程图')

    add_h2('(一) 相关技术简介')
    add_para('本项目主要使用 Python 3.11 进行后端开发，TypeScript + React 19 进行前端开发。'
             '涉及的核心库与模型如下：')

    add_h3('1. 语音识别与音视频处理')
    add_para('faster-whisper：基于 CTranslate2 推理引擎加速的 OpenAI Whisper 实现。'
             '相比原版 transformers 实现，推理速度提升 4 到 8 倍，显存占用降低约 50%。'
             '本项目使用 large-v3 模型，配合 word_timestamps 输出词级时间戳与置信度，'
             '为后续「[?] 置信度可视化」功能提供数据支撑。安装命令：pip install faster-whisper ctranslate2。')
    add_para('ffmpeg：开源音视频处理工具，本项目用于从下载的 mp4 视频中抽取 16 kHz 单声道 WAV 音频，'
             '作为 ASR 输入。同时也用于关键帧抽取（每 10 秒一帧）。')
    add_para('yt-dlp：youtube-dl 的活跃 fork，用于从 B 站、YouTube 下载视频与 metadata（标题、简介、tag）。'
             '项目通过其 cookies 机制访问需登录的高清视频源。')

    add_h3('2. 大语言模型与视觉模型')
    add_para('Transformers + AutoAWQ：HuggingFace transformers 库加载 AWQ 4-bit 量化大模型。'
             'AWQ（Activation-aware Weight Quantization）通过激活值感知的权重量化，'
             '在大幅压缩模型大小的同时保留接近原始精度的能力。'
             '本项目使用 Qwen2.5-7B-Instruct-AWQ，显存占用约 5 GB，可在单张 RTX 4060 上运行。')
    add_para('Qwen2.5-VL-7B-Instruct-AWQ：阿里巴巴开源的视觉语言模型，'
             '能直接接受图片输入并输出自然语言描述。本项目将其用于长视频的关键帧 caption 生成，'
             '作为 LLM 章节切分的视觉线索补充。')
    add_para('Chinese-CLIP：CLIP 的中文复现版本，提供图文跨模态语义对齐的视觉编码器。'
             '本项目用于计算相邻关键帧的视觉距离，配合文本距离做多模态融合。')

    add_h3('3. 文本处理')
    add_para('jieba：中文分词与关键词抽取库。本项目用其 textrank 接口做章节关键词抽取，'
             '也用 cut_for_search 接口做术语表的首次出现检测。安装命令：pip install jieba。')
    add_para('nltk + scikit-learn：TextTiling 实现的依赖。本项目复用 nltk.tokenize.texttiling 模块的'
             'depth score 算法作为兜底章节切分器。')

    add_h3('4. Web 框架与前端')
    add_para('FastAPI + uvicorn：Python 异步 Web 框架，本项目用其暴露 /api/generate、/api/notes、/api/jobs 等接口，'
             '并通过 Server-Sent Events 推送 pipeline 进度。')
    add_para('Next.js 16 + React 19 + TailwindCSS 4：前端使用 Next.js App Router 架构，'
             '支持 SSG 静态化部署。Tailwind CSS 提供原子化样式，配合自实现的 Apple 风格组件库构建用户界面。')
    add_para('Plyr：开源 HTML5 视频播放器，本项目用其封装 React 组件，'
             '通过 currentTime API 实现章节按钮的时间戳跳转。')
    add_para('Framer Motion：React 动画库，用于章节切换、术语高亮等过渡效果。')

    add_h3('5. 报告与展示工具')
    add_para('python-pptx：用于生成课程结题展示 PPT（14 页）。'
             'python-docx：用于生成本实验报告。本项目所有输出材料均通过脚本程序化生成，可复现。')

    add_h2('(二) 流程图')
    add_para('系统整体 pipeline 由 7 个串行模块组成，输入视频 URL 或本地文件，'
             '输出结构化 Markdown 笔记和 Web 可视化页面：')
    add_para('视频下载 → ASR 语音识别 → 视频分类 → LLM 章节切分 → 章标题与摘要生成 → 模板分发 → md / web 双输出',
             first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank()
    add_image_placeholder('[ 截图位 1：系统架构流程图（建议从 PPT 第 3 页系统总览图截屏插入此处） ]')
    add_para('各模块之间通过 JSON 中间产物传递数据（asr.json、chapters.json、summary.json 等），'
             '支持断点续跑：任一中间产物已存在时跳过对应模块。整套 pipeline 全部本地运行，'
             '不依赖任何在线 API 服务，确保数据隐私与可离线复用。')
    add_blank()


# ============================================================
# 第三章
# ============================================================
def chapter_3():
    add_h1('三、数据收集及预处理')

    add_h2('(一) 数据收集')
    add_para('本项目构建了一个跨域 32 视频 benchmark，用于评估系统在不同视频类型上的泛化能力。'
             'Corpus 构成如下：')
    add_bullets([
        '学习类视频：26 个，包括 PPT 教学（王道考研操作系统 / 计算机网络 / 计算机组成原理共 18 个）、'
        '科普视频（线性代数、Python 入门、Claude Code 教程等 8 个）。',
        '实拍类视频：6 个，包括美食 Vlog（日本河豚、可口可乐自助餐等）、AI 资讯播报、AI Agent 教程英文视频等。',
        '时长分布：3 分钟到 25 分钟不等，覆盖短视频与长视频两种场景。',
        '语言分布：中文为主（30 个），英文 2 个（用于双语支持验证）。',
    ])
    add_para('视频源主要来自 B 站，少量来自 YouTube。对每个视频，手工标注了章节边界 gold，'
             '并分别提供 chars chunker (cc=400/800) 和 TextTiling 两套 chunks 数下的标注，'
             '便于后续多角度评估。')
    add_image_placeholder('[ 截图位 2：corpus 概况表或视频列表截图，可从 paper/corpus_status_2026-05-26.md 整理 ]',
                          height_in=2.0)

    add_h2('(二) 数据读取')
    add_para('视频下载使用 yt-dlp 库，自动选择最高清晰度的合并流（视频流 + 音频流）。'
             '对于需要登录的高清视频，通过 cookies 文件机制（data/.cookies/*bilibili*.txt）绕过浏览器 DPAPI 加密限制。'
             '下载完成后调用 ffmpeg 抽取 16 kHz 单声道 WAV 音频作为 ASR 输入。关键代码：')
    add_code([
        '# src/download.py - 下载并抽音频',
        'def download_video(url: str, out_dir: Path):',
        "    ydl_opts = {",
        "        'format': 'bestvideo+bestaudio/best',",
        "        'outtmpl': str(out_dir / '%(id)s.%(ext)s'),",
        "        'cookiefile': str(COOKIES_DIR / 'bilibili.txt'),",
        "        'quiet': True,",
        "    }",
        '    with yt_dlp.YoutubeDL(ydl_opts) as ydl:',
        '        info = ydl.extract_info(url, download=True)',
        "    video_path = out_dir / (info['id'] + '.mp4')",
        "    audio_path = out_dir / (info['id'] + '.wav')",
        "    subprocess.run(['ffmpeg', '-y', '-i', str(video_path),",
        "                    '-ar', '16000', '-ac', '1', str(audio_path)],",
        '                   check=True, capture_output=True)',
        '    return video_path, audio_path',
    ])

    add_h2('(三) 数据合并与清洗（ASR 后处理）')
    add_para('ASR 原始输出存在三类典型问题：专业术语识别错误、卡片回路（连续重复段）、'
             '长视频末尾的 ctranslate2 native abort 崩溃。本项目设计了三层后处理与防御机制：')

    add_h3('1. 术语字典自动注入')
    add_para('从视频 metadata（yt-dlp 抓取的 title + description）中通过 jieba 抽取专业词，'
             '在 ASR initial_prompt 中预置，引导 Whisper 模型偏向正确的同音字。'
             '例如视频标题包含「哲学家进餐问题」，注入字典后 ASR 输出不会再把「哲学家」识别为「哲学者」。'
             '项目还维护了 _GLOBAL（通用专业词）和 computer_org（计算机组成原理域专用词）两层字典，'
             '经 corpus sweep 清理共 668 处 leakage 错误。')

    add_h3('2. LCP 连续重复段去重')
    add_para('Whisper 大模型在 streaming attention 路径下偶发「卡片回路」故障：'
             '同一句话被复读 5 到 50 遍，污染下游的关键词频次统计和章节切分。'
             '本项目使用最长公共前缀（LCP，Longest Common Prefix）算法检测连续相同片段并合并其时间戳：')
    add_code([
        '# src/asr.py - LCP 去重核心',
        'def dedupe_segments(segments, min_run=3):',
        '    # 连续 min_run 段共享公共前缀时合并为一段',
        '    result, run = [], []',
        '    for seg in segments:',
        '        if run and lcp(run[-1].text, seg.text) >= min(',
        '                len(run[-1].text), len(seg.text)) * 0.8:',
        '            run.append(seg)',
        '        else:',
        '            if len(run) >= min_run:',
        '                merged = merge_run(run)',
        '                result.append(merged)',
        '            else:',
        '                result.extend(run)',
        '            run = [seg]',
        '    return result + run',
    ])
    add_para('该机制在王道操作系统 p37 视频（哲学家进餐问题）上将 F1@1 从 0.50 提升到 1.00，'
             '在计算机网络 p38 视频上将严格 F1 从 0.25 提升到 0.75。')

    add_h3('3. 三层 hallucination 防御')
    add_para('ASR 长视频跑到 90% 以上时偶发 ctranslate2 native abort 崩溃，'
             '导致整个识别结果丢失。本项目设计三层防御：'
             '① 解码约束：condition_on_previous_text=False，避免错误累积；'
             '② 1.x 原生 hallucination gate：hallucination_silence_threshold=2.0，'
             'compression_ratio_threshold=2.4；'
             '③ WAL 增量落盘：借鉴数据库 Write-Ahead Logging 思想，每识别完一段就立即写入 partial cache，'
             '崩溃后下次重启自动从最近的 partial 续跑。三层防御实战覆盖 BV1q6（日本河豚 vlog 25 分钟）和'
             'BV1AYR6BsE9U 两个长视频的 native abort，99.99% 覆盖率落盘。')

    add_h2('(四) 特征选择')
    add_para('章节切分阶段需要三类特征：文本特征、视觉特征、视频元数据特征。')
    add_para('文本特征：使用 jieba.analyse.textrank 抽取章节关键词，每段保留 top-K 个；'
             '用 Jaccard 距离计算相邻段的关键词重叠度，作为 TextTiling 的 depth score 信号。')
    add_para('视觉特征：用 ffmpeg 每 10 秒抽一帧关键帧，送入 Chinese-CLIP 视觉编码器得到 512 维特征向量；'
             '计算相邻关键帧的余弦距离，作为视觉切点信号。对长视频额外启用 Qwen2.5-VL-7B-AWQ 视觉语言模型，'
             '对每个 chunk 的代表帧生成自然语言 caption，作为 LLM 章节切分的视觉 prompt cue。')
    add_para('视频元数据特征：标题、描述、tag、时长，用于视频类型分类（教学 / 科普 / Vlog / 演讲）。')

    add_h2('(五) 数据预处理')
    add_para('ASR 转写经过去重清洗后，需要切分为若干 chunks 作为章节切分的输入。本项目支持三种 chunker：')
    add_bullets([
        'chars chunker (cc=400 / cc=800)：按字符数固定切分，简单稳定，作为对比 baseline。',
        'TextTiling chunker：基于 jieba 关键词 Jaccard 距离的 depth score 寻找切点，适合主题切换明显的视频。',
        'LLM chunker：将整段转写一次性喂给 Qwen2.5-7B 输出章节大纲，配合三层兜底使用（详见第四章）。',
    ])
    add_para('各 chunker 输出统一格式为 List[Dict] 包含 start、end、text 三个字段，便于后续模块复用。'
             '预处理后的数据持久化为 chunks.json，供章节切分与摘要生成模块读取。')
    add_blank()


# ============================================================
# 第四章
# ============================================================
def chapter_4():
    add_h1('四、「人工智能」项目的模型化过程')

    add_h2('(一) 完成人工智能技术中模型的选择')
    add_para('本项目涉及多个深度学习模型，每个模型的选型都经过对比实验与权衡：')

    add_h3('1. ASR 模型选择：faster-whisper large-v3')
    add_para('原版 Whisper（transformers 实现）在 25 分钟视频上推理需 15-20 分钟，显存占用约 10 GB。'
             '我们改用 faster-whisper（CTranslate2 加速）+ float16 推理，同一视频仅需 4 分钟，'
             '显存约 5 GB，在性能和资源占用之间取得平衡。large-v3 相对 medium 在中文场景下'
             'WER 降低约 30%，是综合最优选择。')

    add_h3('2. 章节切分模型选择：Qwen2.5-7B-AWQ 替代 Pegasus')
    add_para('项目最初使用 Pegasus（专门的摘要预训练模型）做章节切分与标题生成，'
             '但在 30+ chunks 长视频上发现严重的邻章串台问题（相邻章节标题相互抄袭）。'
             '切换到本地大语言模型 Qwen2.5-7B-AWQ 后，配合 B1 两步法和三层兜底，'
             '在 24 视频 corpus 上达到 100% 切分覆盖率，邻章串台问题彻底解决。'
             '选择 7B 而非 14B 或 32B 模型的核心理由是显存约束：7B AWQ 量化版本仅需 5 GB，'
             '可在消费级 GPU 上运行；同时实测 7B 在章节切分任务上效果接近 14B。')

    add_h3('3. 视觉编码器选择：Chinese-CLIP + Qwen2.5-VL 双层架构')
    add_para('章节切分的视觉信号采用两层架构。短视频（chunks 数 ≤ 15）只用 Chinese-CLIP 的关键帧余弦距离作'
             'tie-breaking；长视频额外启用 Qwen2.5-VL-7B-AWQ 生成关键帧的自然语言 caption，'
             '作为 LLM 切分器的 prompt cue。这一设计的关键发现是：单纯依赖 CLIP 视觉距离会导致 PPT 教学视频过切'
             '（slide 翻页频繁但话题未变），而 VLM caption 能让 LLM 自己判断视觉变化是否对应话题切换。')

    add_h3('4. 视频分类器选择：4 类启发式规则')
    add_para('视频类型分类采用启发式规则而非深度学习模型，原因有三：'
             '① 类别只有 4 类（教学 / 科普 / Vlog / 演讲），规则足以覆盖；'
             '② 启发式规则可解释、可调试，新视频出错时能立即定位失败的判断规则；'
             '③ 在 24 视频上准确率达 24/24，无需引入额外模型成本。'
             '规则基于 ASR 文本特征（专业词密度、口语化程度）+ 视频 metadata（标题关键词、tag、时长）综合判断。')

    add_h2('(二) 完成深度网络的建立')
    add_para('章节切分作为整套系统的核心创新点，其网络架构设计如下：')

    add_h3('1. LLM-as-Segmenter 的 B1 两步法')
    add_para('传统做法是让 LLM 一次性输出「章节起止时间戳 + 章节标题」完整 JSON，'
             '这种「大而全」的输出在 30+ chunks 长视频上经常出错（JSON 格式坏、章节数与 chunks 数不对齐、'
             '标题重复等）。本项目设计 B1 两步法，将切分拆成两次独立调用：')
    add_code([
        '# src/segment_llm.py - B1 两步法',
        'def segment_video(chunks):',
        '    # Step 1: 只让 LLM 输出章节起止 chunk 索引',
        '    outline = call_llm_outline(chunks)',
        '    # Step 2: 对每章独立调用 LLM 起标题',
        '    chapters = []',
        '    for start, end in outline:',
        '        title = call_llm_title(chunks[start:end+1])',
        '        chapters.append(Chapter(start, end, title))',
        '    return chapters',
    ])
    add_para('拆分的好处：'
             '① Step 1 失败可以单独 retry，不会浪费 Step 2 的算力；'
             '② Step 2 只看一章上下文，prompt 更聚焦，标题质量更稳；'
             '③ 错误可被精确定位到是 outline 错了还是 title 错了。')

    add_h3('2. retry-with-feedback 机制')
    add_para('LLM 第一次输出不合格时（JSON 格式错、章节数为 0、章节数超过 chunks 数等），'
             '本项目将具体错误信息写入第二次调用的 user prompt，让 LLM 自我修正。'
             '实测在 24 视频 corpus 上，retry 机制可救回 67% 的初次失败 case。')
    add_code([
        '# src/segment_llm.py - retry with feedback',
        'def segment_with_retry(chunks, max_retries=3):',
        "    feedback = ''",
        '    for attempt in range(max_retries):',
        '        outline = call_llm_outline(chunks, feedback=feedback)',
        '        ok, err = validate_outline(outline, len(chunks))',
        '        if ok:',
        '            return outline',
        "        feedback = '上次输出有问题：' + err + '。请严格按格式重输。'",
        '    raise SegmentationFailed()  # 进入程序化 repair 兜底',
    ])

    add_h3('3. 程序化 _repair_oversize 兜底')
    add_para('retry 仍失败时（约 11% 的 case），进入程序化 repair。'
             '核心规则：若 LLM 输出某章 chunks 数超过总数的 45% 上限，按 chunk 长度等分切两半；'
             '若章节数不足 ceil(duration/6) 下限，按 TextTiling depth score 在最长章内补一刀。'
             'repair 后总能输出合法切分，配合上面的 LLM 主路径，覆盖率达 100%。')

    add_h3('4. 多模态融合：α=0.3 视觉权重')
    add_para('早期 TextTiling 时代，融合公式为：')
    add_para('combined_depth(i) = (1-α) × text_depth(i) + α × visual_depth(i)',
             first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para('在 10 视频 × 2 chunker 的 α sweep 消融实验中发现：')
    add_bullets([
        'α=0.0（纯文本）：实拍 Vlog 上欠切（话题切换镜头明显但文本平稳）。',
        'α=0.3：综合最优甜点，PPT 教学不过切、Vlog 不欠切。',
        'α=0.5：PPT 教学开始过切（slide 翻页被误判为章节切换）。',
        'α=1.0（纯视觉）：所有视频严重过切，不可用。',
    ])
    add_para('当前 LLM 时代视觉信号改作 prompt cue（VLM caption），不再使用 α 加权融合。'
             '但这一历史性消融实验的结论「视觉只能当辅助、不能当主导」指导了后续 VLM 集成的设计。')
    add_image_placeholder('[ 截图位 3：α sweep 折线图（建议从论文 §5.4 数据画一张柱状或折线图插入） ]',
                          height_in=2.5)

    add_h2('(三) 完成模型的训练')
    add_para('本项目所有深度学习模型均使用预训练权重，不进行额外训练。'
             '实际工作量集中在「提示工程」和「知识扩充」两个方面，可视为面向 LLM 的「训练等价物」。')

    add_h3('1. 提示工程：B1 两步法 prompt 设计')
    add_para('Step 1（outline prompt）的核心约束包括：')
    add_bullets([
        '明确输出格式：JSON array of [start_chunk, end_chunk] 二元组。',
        '明确章节数硬约束：max(3, ceil(duration_minutes / 6)) ≤ n_chapters ≤ 总 chunks 数 × 0.6。',
        '明确单章 chunks 数上限：≤ 总 chunks 数 × 0.45，避免 catch-all bias。',
        '明确主题独立性：要求相邻章节主题不重复，禁止使用通用词如「内容」、「部分」。',
    ])
    add_para('Step 2（title prompt）的核心约束包括 ASR snippet 注入（防止 LLM 自由发挥）、'
             '禁止共享前缀（如不允许多章都用「服务程序 X」格式）、'
             '标题字数限制（10-25 字）等。')

    add_h3('2. 字典扩充：术语 corpus 清理')
    add_para('ASR 术语字典在 32 视频 corpus 上跑批后发现 668 处 leakage（错字进入笔记），'
             '通过手工审查 + corpus sweep 共扩充 _GLOBAL 字典 9 个核心术语、'
             'computer_org 域字典 6 个专业词。修复后跨 32 视频 corpus 重跑零 leakage 残留。')

    add_h3('3. 视频分类阈值标定')
    add_para('4 类启发式分类器的核心阈值通过 24 视频标注集合标定，包括：'
             '教学类专业词密度阈值（≥ 1.5%）、Vlog 类口语化指标阈值（「探店」等触发词权重 0.7+）、'
             '演讲类时长下限（≥ 20 分钟）等。最终在 24 视频上准确率 24/24。')

    add_h2('(四) 完成模型的评估')
    add_para('本项目设计了 32 视频跨域 benchmark，从多个角度量化系统效果。')

    add_h3('1. Benchmark 与评估指标')
    add_para('Benchmark 由 26 学习类视频 + 6 实拍类视频构成，每个视频手工标注章节边界 gold，'
             '并分别在 chars chunker 与 TextTiling chunker 下各做一套标注。'
             '评估指标包括：')
    add_bullets([
        'Strict F1：预测边界与 gold 边界完全一致才算正确。适合长视频细粒度评估。',
        'F1@1：允许 ±1 chunk 容差，更适合短视频粗粒度评估，因为短视频章节少，'
        '严格指标的 0/1 二值化问题严重。',
        'LLM 切分覆盖率：LLM 主路径成功率，反映系统稳定性。',
    ])

    add_h3('2. 主结果')
    add_para('关键结果汇总：')
    add_bullets([
        '王道操作系统 p37（哲学家进餐问题）：F1@1 从 0.50 提升到 1.00，靠 ASR 去重救援。',
        '计算机网络 p38（以太网 IEEE 802.3）：严格 F1 从 0.25 提升到 0.75。',
        'LLM 章节切分覆盖率：24 视频 100%（22% 一次过 + 67% retry + 11% 程序救回，'
        'TextTiling fallback 0 次触发）。',
        '视频分类准确率：24/24（教学 / 科普 / Vlog / 演讲 4 类）。',
    ])
    add_image_placeholder('[ 截图位 4：主结果数据柱状图或表格，建议把上面 4 项做成柱状图 ]',
                          height_in=2.5)

    add_h3('3. 消融与跨域')
    add_para('α sweep 消融：在 10 视频 × 2 chunker 上系统化跑 α ∈ {0, 0.1, 0.3, 0.5, 0.7, 1.0}，'
             '确认 α=0.3 是稳健甜点。')
    add_para('多模态 ablation：在 9 视频上对比 txt 路径 vs mm 路径，'
             'mm 路径让 3/9 视频的 LLM attempts 数减少，证明视觉信号确实让大模型「少瞎想」。')
    add_para('跨语言泛化：英文教学视频 EH5jx5qPabU（AI Agents in 25 Minutes）跑通，'
             '适配点包括句号本土化、Qwen 输出英文标题、wrap-up 大小写。')

    add_h2('(五) 完成模型的应用（端到端 Pipeline）')
    add_para('最终的 NoteGen 系统是一个端到端 pipeline，输入视频 URL 或本地文件，'
             '输出 Markdown 笔记和 Web 可视化页面。完整调用链如下：')
    add_code([
        '# src/pipeline.py - 端到端 pipeline',
        'def run_pipeline(url, out_dir):',
        '    # 1. 下载',
        '    video_path, audio_path = download_video(url, out_dir)',
        '    # 2. ASR + 去重 + 三层防御',
        '    asr_result = transcribe(audio_path, terms=extract_terms(url))',
        '    asr_result = dedupe_segments(asr_result)',
        '    # 3. 视频分类',
        '    category = classify_category(asr_result, metadata)',
        '    # 4. LLM 章节切分（B1 两步法 + 三层兜底）',
        '    chunks = chunker.split(asr_result)',
        '    chapters = segment_with_retry(chunks)',
        '    # 5. 章标题 + 摘要 + 章末小结',
        '    summary = summarize(chapters, asr_result, category)',
        '    # 6. 套对应 category 的模板',
        '    md = render_markdown(summary, template=TEMPLATES[category])',
        '    # 7. 写出 md + copy 到 web/public',
        '    save_outputs(md, summary, out_dir)',
    ])
    add_para('整套 pipeline 在 RTX 4060（8 GB 显存）上单视频耗时 5-20 分钟，'
             '视视频时长而定。全程本地运行，无在线 API 依赖。')
    add_blank()


# ============================================================
# 第五章
# ============================================================
def chapter_5():
    add_h1('五、程序可操作性和创新能力情况')

    add_h2('(一) 程序运行效果')
    add_para('系统的最终产物包括 Markdown 笔记文件和 Next.js Web 前端两种呈现形式。'
             '下面通过三个典型案例展示运行效果。')

    add_h3('案例 1：王道操作系统 p37（哲学家进餐问题）')
    add_para('该视频时长 15 分钟，是 ASR 卡片回路救援的典型 case。'
             '原始字幕中「哲学家进餐」段被复读 8 遍，未经 dedupe 处理时章节切分被完全干扰，'
             'F1@1 仅 0.50；经 LCP 去重后 F1@1 提升到 1.00。'
             '下图展示系统输出的「知识点速览」与「章节」区域，可以看到每个知识点都附有'
             '时间戳、缩略图、要点摘要和关键词标签，章节切分清晰准确。')
    if (SCREENSHOT_DIR / '01_os_chapters.png').exists():
        add_image(SCREENSHOT_DIR / '01_os_chapters.png', width_in=4.5,
                  caption='图 5-1  王道 OS p37 知识点速览与章节列表')
    else:
        add_image_placeholder('[ 截图位 5-1：王道 OS p37 知识点速览 ]')

    add_h3('案例 2：日本河豚 Vlog（BV1Q3dHBSEAY）')
    add_para('该视频时长 25 分钟，是 Vlog 域 hallucinate 修复的典型 case。'
             '早期版本中 LLM 给每章生成 abstract 时「自由发挥」，'
             '16 章中有 13 章写出与视频无关的内容（如「心理线」被扩写成「股市心理学」、'
             '「皇上」被写成「清朝历史」）。'
             '修复方案是在 prompt 中注入对应 chunk 的 ASR snippet，强制 LLM 贴合原文。'
             '修复后 3 个 Vlog 视频 21/21 章 0 hallucinate。'
             '下图展示修复后的页面：左侧 Plyr 播放器嵌入视频，'
             '右侧「片段时间线」区域的章节摘要（「日本河豚料理」、「河豚内脏处理」、「河豚市场拍卖」等）'
             '完全贴合视频内容，没有跑题。')
    if (SCREENSHOT_DIR / '02_vlog_abstract.png').exists():
        add_image(SCREENSHOT_DIR / '02_vlog_abstract.png', width_in=6.0,
                  caption='图 5-2  日本河豚 Vlog 修复后的章节摘要')
    else:
        add_image_placeholder('[ 截图位 5-2：vlog 修复后摘要 ]')

    add_h3('案例 3：Web 前端整体效果')
    add_para('前端采用 Next.js 16 + React 19 + Tailwind CSS 4 构建，Apple 风格深色 UI + 流动粒子背景。'
             '页面顶部为 NoteGen logo + 返回首页按钮 + 视频标题；中部为 Plyr 视频播放器；'
             '下方为可点击跳转的章节按钮和当前章节卡片（含进度条与关键词标签）。'
             '整套前端通过 SSG 静态化部署，加载快速、可离线访问。')
    if (SCREENSHOT_DIR / '03_web_overall.png').exists():
        add_image(SCREENSHOT_DIR / '03_web_overall.png', width_in=4.5,
                  caption='图 5-3  Web 前端整体效果')
    else:
        add_image_placeholder('[ 截图位 5-3：Web 前端整体效果 ]')

    add_image_placeholder('[ 截图位 5-4：（可选补充）首页 landing 卡片墙截图，'
                          '或多模态消融 mm-ablation 页面截图 ]', height_in=2.5)

    add_h2('(二) 概念理解及调试')

    add_h3('1. 基本概念理解')
    add_para('本项目涉及多个深度学习概念，开发过程中对以下核心概念有了深入理解：')
    add_bullets([
        '序列到序列建模：Whisper 的 encoder-decoder 架构，以及 LSTM、Transformer 在长序列建模中的差异。',
        '注意力机制：Whisper 和 Qwen2.5 都基于 Transformer 架构，理解 self-attention 与 cross-attention 的工作机制。',
        '量化技术：AWQ（Activation-aware Weight Quantization）将 FP16 模型压缩到 INT4，'
        '在显存与精度间取得平衡，理解了量化感知训练与训练后量化的区别。',
        '多模态对齐：CLIP 通过对比学习实现图文跨模态语义对齐，理解了对比学习与正负样本采样策略。',
        'Prompt Engineering：发现 prompt 设计对 LLM 输出质量影响巨大，结构化 prompt（先 outline 后 title）'
        '比一次性大 prompt 稳定得多。',
    ])

    add_h3('2. 工具使用与调试能力')
    add_para('开发环境使用 VSCode + Cursor IDE，配合 Python venv 虚拟环境管理依赖。'
             '调试技巧包括：')
    add_bullets([
        'ASR 调试：通过 dump asr.json 中间产物查看每段时间戳与文本，结合 word_timestamps 定位识别错误位置。',
        'LLM 调试：在 segment_llm.py 中加入 _debug_dump 函数，将每次 LLM 调用的 input/output 落盘到 logs/，'
        '便于事后复盘失败 case。',
        '前端调试：使用 React DevTools + Next.js 的 SSG 预览模式，在编辑器即时查看修改效果。',
        '性能调试：使用 PyTorch profiler 分析 ASR / LLM / VLM 各阶段耗时，'
        '定位瓶颈在 LLM 章节切分（占总耗时约 40%）。',
    ])

    add_h3('3. 典型调试过程')
    add_para('以「BV1q6 日料探店 13/16 章 abstract 字面 hallucinate」为例，调试流程如下：')
    add_bullets([
        '发现问题：用户反馈「心理线」被写成「股市心理学」，怀疑 LLM 输出乱码。',
        '复现：跑 dryrun 模式，dump 中间 prompt 与 LLM 输出。',
        '定位：发现 prompt 中只传了「章节大纲 + 标题」，未传 ASR snippet。'
        'LLM 在缺乏原始内容上下文时倾向于按标题字面「补全语境」。',
        '修复：在 generate_chapter_abstracts 的 input 中加入对应 chunk.summary snippet，'
        '并在 user_prompt 添加严格约束「abstract 必须基于 snippet 描述实际内容」。',
        '回归：3 vlog corpus（21 章）重跑，0 hallucinate 残留，commit b7c96e8 落地。',
    ])
    add_image_placeholder('[ 截图位 5-5：调试过程示例截图，例如 logs/ 目录里某次 LLM input/output dump，'
                          '或 dedupe 前后字幕对比截图 ]', height_in=2.5)

    add_h2('(三) 创新能力')

    add_h3('1. 算法层创新')
    add_para('（1）LLM-as-Segmenter 的 B1 两步法：'
             '将「切章 + 起标题」拆成两次独立调用，相比传统「一次性大 JSON」方法，'
             '在 24 视频 corpus 上 LLM 主路径成功率从约 30% 提升到 89%。')
    add_para('（2）三层兜底架构（retry-with-feedback → 程序化 repair → TextTiling fallback）：'
             '保证系统在任何输入下都能产出合法的章节切分。'
             '24 视频实测：22% attempt 1 直接通过，67% 由 retry 救回，11% 由 _repair_oversize 救回，'
             'fallback 0 次触发。')
    add_para('（3）LCP 连续重复段去重：'
             '针对 ASR streaming attention 路径下的「卡片回路」故障，'
             '提出基于最长公共前缀的检测与合并算法，并发现「chunker × dedupe」二维耦合效应——'
             'dedupe 本质是「关键词频次去噪」而非简单文本清洗。')
    add_para('（4）多模态融合 α=0.3 甜点：'
             '通过 10 视频 × 2 chunker × α sweep 系统化消融，'
             '在 cross-domain 场景下确定 α=0.3 是稳健甜点（视觉作 tie-breaking，不主导）。')

    add_h3('2. 工程层创新')
    add_para('（1）三层 ASR 健壮性防御：'
             '针对 ctranslate2 在 Windows 平台上的 native abort 故障，'
             '提出「解码约束 + 1.x 原生 hallucination gate + WAL 增量落盘」三层防御，'
             '实战覆盖 BV1q6 和 BV1AYR6BsE9U 两个长 vlog，99.99% 覆盖率落盘。'
             '这是将数据库领域成熟的 WAL 故障兜底范式引入 ASR pipeline 的一次成功尝试。')
    add_para('（2）4 类启发式分类 + 双层模板分发：'
             '用启发式规则在 24 视频上达到 24/24 准确率，无需训练分类模型；'
             '进一步驱动 markdown 与前端模板按 category 差异化渲染（教学用知识点速览，Vlog 用探店摘要）。')
    add_para('（3）Stress Test 范式：'
             'v6 阶段通过 scripts/stress_test_* 对系统兜底机制做主动 stress test，'
             '暴露出「野外永远不触发因此被误判为不需要」的隐藏 bug（auto_subs n≥4 broken），'
             '总结出「对每个兜底必须配 stress test」的工程方法论。')

    add_h3('3. 产品层创新')
    add_para('（1）学习场景 5 件套 markdown 模板：'
             '在传统「章节 + 段落」两级结构上注入顶部摘要卡、HTML 锚点 TOC、知识点速览、'
             '术语表（含首次出现链接 + 上下文 snippet）、抽取式章末小结，'
             '提供「二次学习」的快捷入口。')
    add_para('（2）中英双语 toggle：'
             '前端实现章节标题、章节摘要、关键词的中英文同步切换；'
             '后端通过 Qwen2.5 翻译自动生成 _zh / _en 双字段。')
    add_para('（3）VLM caption 视觉信号升级：'
             '用 Qwen2.5-VL-7B-AWQ 生成关键帧自然语言 caption 作为 LLM 切分的 prompt cue，'
             '比单纯 CLIP 距离更稳定，对 PPT 教学过切风险有显著缓解。')
    add_blank()


# ============================================================
# 第六章
# ============================================================
def chapter_6():
    add_h1('六、结束语')

    add_h3('方法优缺点总结')
    add_para('本项目所采用的「本地 LLM-as-Segmenter + 三层兜底 + 学习场景模板分发」方案具有明显优点：'
             '首先，章节切分覆盖率达 100%，系统在任何输入下都能产出合法笔记，无失败 case；'
             '其次，全本地部署不依赖在线 API，保护用户数据隐私的同时降低运行成本；'
             '再次，4 类视频分类的差异化模板分发解决了通用产品「一刀切」的结构空洞问题；'
             '最后，三层 ASR 健壮性防御对 ctranslate2 native abort 等隐蔽故障有可靠的兜底。')
    add_para('但本方案也存在若干局限性。'
             '一是对硬件有一定要求：完整 pipeline 需要至少 5 GB 显存的 GPU，'
             '对于无独立显卡的用户不友好；'
             '二是 Qwen2.5-7B 在 30+ chunks 极长视频上偶发失效（如 p93 单一案例的 Ch1 LLM 幻觉走样），'
             '需要更大模型或更细粒度的章节切分策略；'
             '三是 PPT slide cue 检测在跨视频间的阈值标定困难，目前默认禁用此模块；'
             '四是当前 benchmark 规模（32 视频）相对学术界主流数据集偏小，'
             '结论的普适性需要更大规模验证。')

    add_h3('错误与解决方法')
    add_para('项目开发过程中遇到的主要错误及解决方法包括：')
    add_bullets([
        'ASR ctranslate2 native abort：通过 WAL 增量落盘 + hallucination gate 三层防御解决。',
        'LLM abstract 字面 hallucinate：通过在 prompt 注入 ASR snippet 解决。',
        '章标题「服务程序 X」共享前缀降级：通过 snippet 喂 LLM + prompt 共享前缀禁令 + Python 兜底 + '
        '0.7 视频主题守门解决（J7 三件套）。',
        'ASR mask [?] 字符泄漏到用户端：增加 _strip_qmask 与 post-fix 脚本批量清理。',
        'autoawq 在 Windows 上 DLL 注入失败：通过 transformers 4.49 + autoawq 0.2.6 + shim 提模块顶层解决。',
        'GPU 黑屏花屏事故：撤销大模型并行常驻方案，定红线「所有大模型必须串行加载，用完释放」。',
    ])

    add_h3('改进与拓展')
    add_para('未来可在以下方向继续改进和拓展：')
    add_bullets([
        '扩展视频类型：目前主要覆盖教学、科普、Vlog、演讲 4 类，未来可扩展到讲座、播客、纪录片等更多场景。',
        '更小更快的章节切分模型：探索 Qwen2.5-3B / 1.5B 等更小模型，'
        '在准确率与推理速度之间寻找新的平衡点。',
        '大规模用户主观打分评估：当前评估以客观指标（F1、覆盖率）为主，'
        '可补充用户主观打分（笔记可读性、术语表完整度等）的大规模问卷调查。',
        '实时流式 pipeline：当前是离线 batch 模式，'
        '未来可改造为边录边处理的实时流式系统，适配直播场景。',
        '知识图谱集成：将术语表升级为跨视频的知识图谱，'
        '支持「从一个视频跳转到另一个视频相同概念」的关联学习。',
        '云端部署与多用户：当前是单机本地版本，可改造为多租户云服务，'
        '结合容器化部署支持团队/班级级别的共享笔记。',
    ])
    add_para('通过本项目的开发，我系统性地学习了深度学习在视频理解、自然语言处理、'
             '多模态融合三个方向的实际应用，也对工程化部署、调试方法论、'
             '用户体验设计有了更深刻的认识。'
             '项目所采用的「算法 + 工程」双重创新理念，'
             '尤其是将数据库 WAL 思想引入 ASR pipeline、'
             '将启发式规则用于代替深度学习分类器等设计取舍，'
             '为后续类似工程化 AI 应用提供了可借鉴的方法论。')
    add_blank()


# ============================================================
# 参考文献
# ============================================================
def references():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('参考文献')
    set_cn_font(r, font='黑体', size=16, bold=True)
    add_blank()
    refs = [
        '[1] Hearst, M. A. TextTiling: Segmenting text into multi-paragraph subtopic passages. '
        'Computational Linguistics, 1997, 23(1): 33-64.',
        '[2] Radford, A., Kim, J. W., Hallacy, C., et al. Learning transferable visual models from natural language supervision. '
        'In ICML, 2021: 8748-8763.',
        '[3] Yang, A., Pan, J., Lin, J., et al. Chinese CLIP: Contrastive vision-language pretraining in Chinese. '
        'arXiv preprint arXiv:2211.01335, 2022.',
        '[4] Radford, A., Kim, J. W., Xu, T., et al. Robust speech recognition via large-scale weak supervision. '
        'In ICML, 2023: 28492-28518. (OpenAI Whisper)',
        '[5] Bai, J., Bai, S., Yang, S., et al. Qwen technical report. '
        'arXiv preprint arXiv:2309.16609, 2023.',
        '[6] Wang, P., Bai, S., Tan, S., et al. Qwen2-VL: Enhancing vision-language model perception of the world at any resolution. '
        'arXiv preprint arXiv:2409.12191, 2024.',
        '[7] Lin, J., Tang, J., Tang, H., et al. AWQ: Activation-aware weight quantization for LLM compression and acceleration. '
        'In MLSys, 2024.',
        '[8] Koshorek, O., Cohen, A., Mor, N., et al. Text segmentation as a supervised learning task. '
        'In NAACL-HLT, 2018: 469-473.',
        '[9] Zhang, J., Zhao, Y., Saleh, M., et al. PEGASUS: Pre-training with extracted gap-sentences for abstractive summarization. '
        'In ICML, 2020: 11328-11339.',
        '[10] Lewis, M., Liu, Y., Goyal, N., et al. BART: Denoising sequence-to-sequence pre-training for natural language generation, translation, and comprehension. '
        'In ACL, 2020: 7871-7880.',
        '[11] OpenNMT. CTranslate2: Fast inference engine for Transformer models. '
        'https://github.com/OpenNMT/CTranslate2, 2025.',
        '[12] SYSTRAN. faster-whisper: Faster Whisper transcription with CTranslate2. '
        'https://github.com/SYSTRAN/faster-whisper, 2025.',
        '[13] yt-dlp project. yt-dlp: A youtube-dl fork with additional features and fixes. '
        'https://github.com/yt-dlp/yt-dlp, 2025.',
        '[14] Vercel. Next.js 16 documentation. https://nextjs.org/docs, 2026.',
        '[15] Goodfellow, I., Bengio, Y., Courville, A. Deep Learning. MIT Press, 2016.',
        '[16] 周志华. 机器学习. 清华大学出版社, 2016.',
        '[17] 李航. 统计学习方法（第 2 版）. 清华大学出版社, 2019.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(20)
        r = p.add_run(ref)
        set_cn_font(r, font='宋体', size=12)
    page_break()


# ============================================================
# 成绩、评语、教师签名
# ============================================================
def grading_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('成绩、评语、教师签名')
    set_cn_font(r, font='黑体', size=16, bold=True)
    add_blank()
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = '成绩'
    t.rows[0].cells[1].text = ''
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, font='宋体', size=12, bold=True)
    add_blank()
    t2 = doc.add_table(rows=1, cols=1)
    t2.style = 'Table Grid'
    cell = t2.rows[0].cells[0]
    cell.text = ''
    for _ in range(6):
        cell.add_paragraph()
    p = cell.add_paragraph('评语：')
    for r in p.runs:
        set_cn_font(r, font='宋体', size=12)
    for _ in range(4):
        cell.add_paragraph()
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('教师签名：                       ')
    set_cn_font(r, font='宋体', size=12)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('2026  年   6   月   10   日')
    set_cn_font(r, font='宋体', size=12)


# ============================================================
# 主流程
# ============================================================
build_cover()
build_toc()
chapter_1()
chapter_2()
chapter_3()
chapter_4()
chapter_5()
chapter_6()
references()
grading_page()

doc.save(str(OUT))
print('OK:', OUT)
print('  paragraphs:', len(doc.paragraphs))
print('  tables:', len(doc.tables))
