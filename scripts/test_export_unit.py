"""导出 docx 断言：md_to_docx 转换正确性 + POST /api/export/docx 端点
（空 body 400 / 超限 413 / 正常 200 返回合法 docx + 中文文件名 RFC5987）。
临时库；不需 Redis / GPU。
Run: PYTHONIOENCODING=utf-8 .venv/Scripts/python.exe scripts/test_export_unit.py"""
import io
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))  # repo root: 找到 server.py
os.environ["NOTEGEN_DB_PATH"] = os.path.join(tempfile.mkdtemp(), "t.db")
import db  # noqa: E402
db.set_db_path(os.environ["NOTEGEN_DB_PATH"])

from fastapi.testclient import TestClient  # noqa: E402
import server  # noqa: E402
import md_to_docx  # noqa: E402
from docx import Document  # noqa: E402

FAILS = []
def check(cond, msg):
    print(("PASS" if cond else "FAIL") + " : " + msg)
    if not cond:
        FAILS.append(msg)


# ============ (a) md_to_docx.convert：本项目导出 md 实际用到的元素 ============
SAMPLE_MD = """# 操作系统第一讲

> 来源: https://www.bilibili.com/video/BV1xxx
> UP 主: 某老师

## 本视频讲了什么

进程与线程的区别，**调度**算法概览。

**你将学到**

- 进程的定义
- `fork()` 的语义

## 1. 进程基础（00:00 - 12:30）

进程是资源分配的最小单位。

- [00:15] 进程定义
- [03:42] PCB 结构
"""

tmp = Path(tempfile.mkdtemp())
md_p = tmp / "in.md"
docx_p = tmp / "out.docx"
md_p.write_text(SAMPLE_MD, encoding="utf-8")
md_to_docx.convert(md_p, docx_p)
check(docx_p.is_file() and docx_p.stat().st_size > 0, "convert 产出非空 .docx")

doc = Document(str(docx_p))
texts = [p.text for p in doc.paragraphs]
all_text = "\n".join(texts)
check("操作系统第一讲" in all_text, "H1 标题进入文档")
check("进程基础" in all_text, "章节标题进入文档")
check("进程的定义" in all_text, "bullet 列表项进入文档")
check("[00:15] 进程定义" in all_text, "时间戳行进入文档")
check("调度" in all_text, "加粗 inline 文本保留")

heads = [p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")]
check(len(heads) >= 3, f"标题层级映射为 Word Heading（{len(heads)} 个）")


# ============ (b) POST /api/export/docx 端点 ============
c = TestClient(server.app)

r = c.post("/api/export/docx", json={"markdown": "", "filename": "x"})
check(r.status_code == 400, f"空 markdown → 400（{r.status_code}）")

r = c.post("/api/export/docx",
           json={"markdown": "a" * 2_000_001, "filename": "x"})
check(r.status_code == 413, f"超过 2MB → 413（{r.status_code}）")

r = c.post("/api/export/docx",
           json={"markdown": SAMPLE_MD, "filename": "操作系统第一讲.docx"})
check(r.status_code == 200, f"正常导出 → 200（{r.status_code}）")
ct = r.headers.get("content-type", "")
check("wordprocessingml" in ct, f"Content-Type 是 docx（{ct}）")
cd = r.headers.get("content-disposition", "")
check("filename*=UTF-8''" in cd and "%E6%93%8D%E4%BD%9C" in cd,
      "Content-Disposition 带 UTF-8 中文文件名")
doc2 = Document(io.BytesIO(r.content))
body2 = "\n".join(p.text for p in doc2.paragraphs)
check("操作系统第一讲" in body2, "返回字节可被 python-docx 打开且含标题")

# 文件名兜底：无扩展名/危险字符
check(server._safe_export_filename(None) == "notes.docx", "空文件名 → notes.docx")
check(server._safe_export_filename("a/b\\c:d") == "abcd.docx", "路径分隔符被剥离")
check(server._safe_export_filename("讲义") == "讲义.docx", "自动补 .docx")

print()
if FAILS:
    print(f"{len(FAILS)} FAILED")
    sys.exit(1)
print("ALL PASS")
