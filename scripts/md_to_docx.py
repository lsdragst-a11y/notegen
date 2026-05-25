"""把 paper/draft.md 转 Word（.docx）。

只处理本项目实际用到的 markdown 元素：
  # / ## / ### / #### / ##### 标题
  普通段落 / blockquote (>)
  bullet 列表 (- / *)
  numbered 列表 (1. / 2.)
  code fence ```...```
  inline code `...`, bold **...**, italic *...*（最常用）
  table | a | b | + |---|---| 分隔
  horizontal rule ---

不依赖 pandoc。单文件运行：
  .venv/Scripts/python.exe scripts/md_to_docx.py paper/draft.md paper/draft.docx
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


HEADING_PT = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11}
INLINE_CODE_FONT = "Consolas"
BODY_FONT = "Microsoft YaHei"  # 默认中文字体
MONO_FONT = "Consolas"


def _set_run_default(run, size_pt=11, font=BODY_FONT):
    run.font.name = font
    run.font.size = Pt(size_pt)
    # 中文字体兼容 (W3C 标准要 east_asia)
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def _add_inline(paragraph, text: str, size_pt: int = 11):
    """处理 inline: **bold**, *italic*, `code`, [link](url)"""
    # 复合 regex 一次切：(**...**) | (*...*) | (`...`) | ([text](url)) | plain
    pattern = re.compile(
        r"(\*\*[^*\n]+\*\*)|(\*[^*\n]+\*)|(`[^`\n]+`)|(\[[^\]\n]+\]\([^)\n]+\))"
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            r = paragraph.add_run(text[last:m.start()])
            _set_run_default(r, size_pt)
        s = m.group(0)
        if s.startswith("**"):
            r = paragraph.add_run(s[2:-2])
            r.bold = True
            _set_run_default(r, size_pt)
        elif s.startswith("`"):
            r = paragraph.add_run(s[1:-1])
            r.font.name = MONO_FONT
            r.font.size = Pt(size_pt - 1)
        elif s.startswith("["):
            # [text](url) — 简化处理：显示 text + (url) 后缀
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", s)
            if link_m:
                r = paragraph.add_run(link_m.group(1))
                r.underline = True
                r.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
                _set_run_default(r, size_pt)
        elif s.startswith("*"):
            r = paragraph.add_run(s[1:-1])
            r.italic = True
            _set_run_default(r, size_pt)
        last = m.end()
    if last < len(text):
        r = paragraph.add_run(text[last:])
        _set_run_default(r, size_pt)


def _add_code_block(doc, code: str, lang: str = ""):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = MONO_FONT
    r.font.size = Pt(9)
    # 浅灰底色（用 paragraph shading）
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def _add_table(doc, rows: list[list[str]]):
    if not rows or not rows[0]:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= len(tbl.rows[ri].cells):
                continue
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline(p, cell_text.strip(), size_pt=10)
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    # 设默认字体
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)

    i = 0
    n = len(lines)
    in_code = False
    code_buf: list[str] = []
    code_lang = ""

    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # code fence
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_buf = []
            else:
                _add_code_block(doc, "\n".join(code_buf), code_lang)
                in_code = False
                code_buf = []
                code_lang = ""
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # heading
        h_match = re.match(r"^(#{1,5})\s+(.+?)\s*#*\s*$", stripped)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2)
            p = doc.add_heading(level=level)
            _add_inline(p, heading_text, size_pt=HEADING_PT.get(level, 11))
            for r in p.runs:
                r.bold = True
            i += 1
            continue

        # horizontal rule
        if re.match(r"^---+\s*$", stripped) or re.match(r"^\*\*\*+\s*$", stripped):
            p = doc.add_paragraph()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "808080")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # table (detect | header | + |---|)
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s\-:|]+\|\s*$", lines[i + 1].rstrip()):
            # gather table rows
            rows = []
            while i < n and lines[i].rstrip().startswith("|"):
                row_line = lines[i].rstrip()
                if re.match(r"^\|[\s\-:|]+\|\s*$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                rows.append(cells)
                i += 1
            _add_table(doc, rows)
            continue

        # blockquote
        if stripped.startswith(">"):
            q_lines = []
            while i < n and lines[i].rstrip().startswith(">"):
                q_lines.append(lines[i].rstrip().lstrip(">").lstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            _add_inline(p, " ".join(q_lines))
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        # bullet list
        if re.match(r"^[\-*]\s+", stripped):
            li_text = re.sub(r"^[\-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, li_text)
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", stripped):
            li_text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, li_text)
            i += 1
            continue

        # empty line
        if not stripped:
            i += 1
            continue

        # normal paragraph — collect contiguous non-special lines
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("```") or
                nxt.startswith("|") or nxt.startswith(">") or
                re.match(r"^[\-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt) or
                re.match(r"^---+\s*$", nxt)):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        _add_inline(p, " ".join(para_lines))

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"OK → {docx_path}", flush=True)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
